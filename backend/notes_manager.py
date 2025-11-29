from pathlib import Path
from datetime import datetime
from docx import Document
from config import BASE_DIR
from backend.encryption_utils import generate_random_key, encrypt_file, decrypt_file

NOTES_DIR = BASE_DIR / "data" / "notes"

def ensure_notes_dir():
    NOTES_DIR.mkdir(parents=True, exist_ok=True)


def save_note(title: str, content: str):
    ensure_notes_dir()

    safe = "".join(c for c in title if c.isalnum() or c in (" ", "_", "-")).rstrip()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"{safe.replace(' ', '_')}_{timestamp}.docx"
    file_path = NOTES_DIR / file_name


    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.save(file_path)

    
    key = generate_random_key()
    encrypt_file(file_path, key)

    return file_name, key


def list_notes():
    ensure_notes_dir()
    notes = [f.name for f in NOTES_DIR.glob("*.docx")]
    return sorted(notes, reverse=True)


def load_note_content(filename: str, user_key: str):
    ensure_notes_dir()
    file_path = NOTES_DIR / filename

    if not file_path.exists():
        return None, None

    ok = decrypt_file(file_path, user_key)
    if not ok:
        return None, None

    doc = Document(file_path)
    title = doc.paragraphs[0].text if doc.paragraphs else filename
    content = "\n".join(p.text for p in doc.paragraphs[1:])
    encrypt_file(file_path, user_key)

    return title, content


def update_note(filename: str, current_key: str, new_title: str, new_content: str):
    ensure_notes_dir()
    file_path = NOTES_DIR / filename

    if not file_path.exists():
        raise FileNotFoundError(f"Note not found: {file_path}")

    ok = decrypt_file(file_path, current_key)
    if not ok:
        raise ValueError("Incorrect current key")

    doc = Document()
    doc.add_heading(new_title, level=1)
    doc.add_paragraph(new_content)
    doc.save(file_path)

    new_key = generate_random_key()
    encrypt_file(file_path, new_key)

    return new_key

def count_notes():
    ensure_notes_dir()
    return len(list(NOTES_DIR.glob("*.docx")))
